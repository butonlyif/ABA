"""
商业计划书（Word 详版）
定位：面向自闭症儿童家长的 AI 包裹式支撑平台
立论：一个自闭症孩子冲击整个家庭系统 —— 市场只治孩子，没人接住家长
用途：对外融资 + 政府/公益合作
配色：ABA 蓝 / 教练青绿（与产品一致）
说明：正文强调一律用「」，避免与 Python 字符串定界符冲突。
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ---------- 颜色 ----------
C_ABA      = RGBColor(0x3A, 0x7B, 0xC8)
C_ABA_DARK = RGBColor(0x1F, 0x4E, 0x8C)
C_COACH    = RGBColor(0x0F, 0x76, 0x6E)
C_ACCENT   = RGBColor(0xE8, 0x7A, 0x4E)
C_INK      = RGBColor(0x1F, 0x29, 0x37)
C_TEXT     = RGBColor(0x37, 0x41, 0x51)
C_MUTE     = RGBColor(0x6B, 0x72, 0x80)
C_LINE     = RGBColor(0xE5, 0xE7, 0xEB)
C_ZEBRA    = RGBColor(0xF7, 0xFA, 0xFC)
C_TIPBG    = RGBColor(0xEC, 0xF4, 0xFB)
C_DANGER   = RGBColor(0xB9, 0x1C, 0x1C)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

DIAGRAM_DIR = "/Users/wangxin/Documents/bjcontents/AI_codex/diagrams/"


def _hex(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

def cell_bg(cell, c):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), _hex(c))
    tcPr.append(shd)

def cell_borders(cell, c=C_LINE, sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for e in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{e}'); el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz); el.set(qn('w:color'), _hex(c)); b.append(el)
    tcPr.append(b)

def cell_no_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for e in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{e}'); el.set(qn('w:val'), 'nil'); b.append(el)
    tcPr.append(b)

def para_shade(p, c):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), _hex(c)); pPr.append(shd)

def para_left_bar(p, c, sz='28'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '8'); left.set(qn('w:color'), _hex(c))
    pBdr.append(left); pPr.append(pBdr)

CJK_FONT = 'Noto Sans SC'

def font(run, size=11, color=C_TEXT, bold=False, italic=False, name=CJK_FONT):
    run.font.size = Pt(size); run.font.color.rgb = color
    run.bold = bold; run.italic = italic
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name); rFonts.set(qn('w:eastAsia'), name)

# ---------- 构件 ----------

def chapter_banner(doc, num, title, accent=C_ABA):
    tbl = doc.add_table(rows=1, cols=2); tbl.autofit = False
    tbl.columns[0].width = Cm(2.2); tbl.columns[1].width = Cm(14.0)
    c0 = tbl.rows[0].cells[0]; c0.width = Cm(2.2)
    cell_bg(c0, accent); c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p0.add_run(num), 26, WHITE, True)
    c1 = tbl.rows[0].cells[1]; c1.width = Cm(14.0)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = c1.paragraphs[0]; p1.paragraph_format.left_indent = Cm(0.4)
    font(p1.add_run(title), 20, C_INK, True)
    for cell in (c0, c1): cell_no_border(cell)
    pl = doc.add_paragraph()
    pl.paragraph_format.space_before = Pt(4); pl.paragraph_format.space_after = Pt(10)
    pPr = pl._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), _hex(accent))
    pBdr.append(bottom); pPr.append(pBdr)

def section(doc, text, accent=C_ABA):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    font(p.add_run("■  "), 12, accent, True)
    font(p.add_run(text), 14, C_INK, True)
    return p

def subhead(doc, text, accent=C_ABA):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    font(p.add_run(text), 11.5, accent, True)
    return p

def body(doc, text, size=10.5, lead=1.45):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = lead; p.paragraph_format.space_after = Pt(5)
    font(p.add_run(text), size, C_TEXT)
    return p

def bullet(doc, text, accent=C_ABA, indent=0.8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    font(p.add_run("• "), 10.5, accent, True)
    font(p.add_run(text), 10.5, C_TEXT)
    return p

def lead_bullet(doc, head, text, accent=C_ABA, indent=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    font(p.add_run("• "), 10.5, accent, True)
    font(p.add_run(head + "  "), 10.5, C_INK, True)
    font(p.add_run(text), 10.5, C_TEXT)
    return p

def numbered(doc, n, text, accent=C_ABA):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.space_after = Pt(2)
    font(p.add_run(f"{n:>2}.  "), 10.5, accent, True)
    font(p.add_run(text), 10.5, C_TEXT)
    return p

def callout(doc, text, kind="tip"):
    accent = {"tip": C_ABA, "note": C_COACH, "danger": C_DANGER}.get(kind, C_ABA)
    label = {"tip": "洞察", "note": "提示", "danger": "重要"}.get(kind, "洞察")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.2); p.paragraph_format.line_spacing = 1.4
    para_left_bar(p, accent)
    para_shade(p, C_TIPBG if kind != "danger" else RGBColor(0xFD, 0xF2, 0xF2))
    font(p.add_run(f" {label}  "), 9.5, accent, True)
    font(p.add_run(text), 10, C_TEXT)
    return p

def table(doc, headers, data, widths=None, accent=C_ABA):
    t = doc.add_table(rows=len(data) + 1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell_bg(cell, accent); cell_borders(cell, accent)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        font(p.add_run(h), 10, WHITE, True)
    for ri, row in enumerate(data):
        bg = C_ZEBRA if ri % 2 == 0 else WHITE
        for ci, txt in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell_bg(cell, bg); cell_borders(cell, C_LINE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            font(p.add_run(str(txt)), 9.5, C_TEXT)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells): row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def metric_row(doc, metrics):
    t = doc.add_table(rows=1, cols=len(metrics)); t.autofit = False
    w = 16.2 / len(metrics)
    for i, (big, small, color) in enumerate(metrics):
        cell = t.rows[0].cells[i]; cell.width = Cm(w)
        cell_bg(cell, C_ZEBRA); cell_no_border(cell)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        font(p.add_run(big), 18, color, True)
        p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        font(p2.add_run(small), 9, C_MUTE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def diagram(doc, filename, caption="", width=15.0):
    path = os.path.join(DIAGRAM_DIR, filename)
    if not os.path.exists(path):
        print("缺图：" + filename); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width))
    if caption:
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(10)
        font(p2.add_run(caption), 9, C_MUTE, italic=True)

def _page_field(p):
    run = p.add_run()
    a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
    b = OxmlElement('w:instrText'); b.set(qn('xml:space'), 'preserve'); b.text = 'PAGE'
    c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'end')
    run._r.append(a); run._r.append(b); run._r.append(c)
    font(run, 9, C_MUTE)

def setup_footer(doc):
    sec = doc.sections[0]; sec.different_first_page_header_footer = True
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(fp.add_run("星星家庭支撑平台 · 商业计划书      第 "), 9, C_MUTE)
    _page_field(fp); font(fp.add_run(" 页"), 9, C_MUTE)


# ============================================================
def build():
    doc = Document()
    _ns = doc.styles['Normal']; _ns.font.name = CJK_FONT
    _rf = _ns.element.get_or_add_rPr().get_or_add_rFonts()
    _rf.set(qn('w:ascii'), CJK_FONT); _rf.set(qn('w:hAnsi'), CJK_FONT); _rf.set(qn('w:eastAsia'), CJK_FONT)
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.6); s.right_margin = Cm(2.6)
    setup_footer(doc)

    # ---------- 封面 ----------
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    cov = os.path.join(DIAGRAM_DIR, "bp_cover.png")
    if os.path.exists(cov):
        p.add_run().add_picture(cov, width=Cm(13.5))
    doc.add_page_break()

    # ---------- 执行摘要 ----------
    pt = doc.add_paragraph()
    font(pt.add_run("执行摘要"), 20, C_ABA_DARK, True)
    pt.paragraph_format.space_after = Pt(4)
    ps = doc.add_paragraph(); para_left_bar(ps, C_ABA)
    font(ps.add_run(" EXECUTIVE SUMMARY"), 10, C_MUTE, True)
    ps.paragraph_format.space_after = Pt(12)

    body(doc,
        "中国有超过 1300 万孤独症（自闭症）患者，其中 0–14 岁儿童 200–300 万，并以每年约 20 万的速度增长。"
        "围绕这些孩子，已经形成了千亿级的康复市场——但几乎所有的产品和机构，都只服务「孩子的康复」。"
        "在孩子被干预的同时，承担一切的家长，长期处于经济重压、照护耗竭、心理危机、婚姻与社交困境之中，却几乎无人接住。")
    body(doc,
        "我们做的，是一套面向自闭症儿童家长的 AI「包裹式」支撑平台。它由两条线组成：「帮孩子」的 ABA 智能助手"
        "（科学干预、训练记录、数据看板），与「顾家长」的人生教练（基于 ACT 接纳承诺疗法的情绪支持、成长路径与"
        "24 小时 AI 陪伴）。两者同账号、数据互通——孩子的进步会反哺对家长的理解，家长状态的改善又会提升干预质量。")
    callout(doc,
        "我们不是又一个「教孩子」的工具。我们是第一个把「整个家庭」作为服务对象、用 AI 规模化托住家长的平台。", "tip")

    subhead(doc, "三个为什么", C_ABA)
    lead_bullet(doc, "为什么是这件事——", "刚需且未被满足：孩子的康复有人做，家长的崩溃没人管，而家长才是干预能否持续的决定因素。")
    lead_bullet(doc, "为什么是现在——", "政策（七部门 2024–2028 关爱行动明确「家庭暖心」「政府购买心理疏导/喘息服务」）+ AI（大模型让一对一心理陪伴的边际成本趋近于零）双重拐点。")
    lead_bullet(doc, "为什么是我们——", "已有可运行的双 App 产品（ABA 主应用 + 人生教练），具备 ACT 引擎、安全层与 SSO 打通，可立即试点。")

    metric_row(doc, [
        ("1300万+", "孤独症患者", C_ABA),
        ("1800亿", "年康复市场", C_COACH),
        ("20万/年", "新增患者", C_ACCENT),
        ("0", "系统服务家长的对手", C_ABA_DARK),
    ])
    doc.add_page_break()

    # ---------- 一 ----------
    chapter_banner(doc, "一", "行业背景：被忽视的「整个家庭」", C_ABA)

    section(doc, "1.1  一个庞大且持续增长的群体", C_ABA)
    body(doc,
        "据中国残联 2023 年普查数据，我国孤独症患者已超 1300 万人，并以每年近 20 万人的速度增长，"
        "发病率已居各类精神残疾首位；国家卫健委数据显示儿童患病率约 1/143（约 0.7%），0–14 岁儿童患者达 200–300 万。"
        "这意味着背后是数百万个被深刻改变的家庭。")
    table(doc,
        headers=["维度", "数据", "来源"],
        data=[
            ["患者总数", "超过 1300 万人", "中国残联 2023 普查报告"],
            ["儿童患者", "0–14 岁约 200–300 万人", "行业康复报告"],
            ["新增速度", "约 20 万人 / 年", "中国残联"],
            ["儿童患病率", "约 1/143（≈0.7%）", "国家卫健委 2022"],
            ["性别比例", "男:女 约 4:1", "多来源一致"],
        ],
        widths=[3.5, 5.5, 7])

    section(doc, "1.2  真正的痛点：一个孩子，冲击整个家庭系统", C_ABA)
    body(doc,
        "自闭症不是一个孩子的病，而是一个家庭长期、全面的压力源。在以孩子为中心的干预之外，"
        "家长承受着层层叠加的冲击——而这些，几乎从未被任何产品系统性地回应。")
    diagram(doc, "bp_family_impact.png", "图 1  一个自闭症孩子对家庭系统的多维冲击")
    table(doc,
        headers=["冲击维度", "家庭实际处境"],
        data=[
            ["经济重担", "家庭年均康复支出常达 5–6 万元，且往往需要父母一方辞职全程陪训，收入与支出双向恶化。"],
            ["照护耗竭", "有效干预理想标准为每周 20–40 小时，照护几乎全年无休，身心长期透支。"],
            ["心理危机", "家长焦虑、抑郁等心理问题检出率显著高于普通人群，却长期缺乏专业、可负担的支持。"],
            ["婚姻压力", "对干预方式、责任分担的分歧，使家庭冲突与离异风险明显上升。"],
            ["社交孤立", "因孩子的特殊行为被误解、回避社交，原有支持网络逐步萎缩。"],
            ["职业中断", "陪训导致晋升停滞、收入下降，重返职场困难。"],
        ],
        widths=[3, 13])
    callout(doc,
        "家长是干预的「第一执行者」——家长一旦垮掉，再科学的训练方案也无法落地。"
        "支撑家长，本质上就是在保护孩子的康复成效。这是我们立项的核心逻辑。", "note")

    section(doc, "1.3  供给侧的结构性缺口", C_ABA)
    body(doc, "即使只看「孩子」这一侧，供给也远未满足需求；面向「家长」的供给则几乎是空白。")
    table(doc,
        headers=["缺口", "现状", "含义"],
        data=[
            ["康复师缺口", "在册康复师约 2.3 万人，按 1:10 服务比缺口约 30 万人", "线下产能无法覆盖需求"],
            ["干预时长不足", "超 70% 的 0–6 岁患儿每周机构干预 < 12 小时（理想 20–40h）", "家庭场景必须补位"],
            ["优质资源稀缺", "仅约 30% 医生认为本地有满足需求的优质康复资源", "资源分布极不均衡"],
            ["家长支持", "几乎没有面向家长心理与能力的系统化产品", "蓝海，且是我们的主战场"],
        ],
        widths=[3, 9, 4])
    doc.add_page_break()

    # ---------- 二 ----------
    chapter_banner(doc, "二", "市场分析", C_ABA)

    section(doc, "2.1  市场规模", C_ABA)
    body(doc,
        "按家庭年均康复花费约 6 万元测算，0–6 岁小龄康复市场约 600 亿元/年，0–18 岁孤独症康复市场约 1800 亿元/年；"
        "若叠加注意力障碍、言语障碍，0–18 岁三项发展障碍康复市场规模达约 9600 亿元/年"
        "（北大医疗脑健康《2020 年度儿童发展障碍康复报告》）。")
    diagram(doc, "bp_market.png", "图 2  中国康复市场规模与患者 / 人才缺口", width=15.5)

    section(doc, "2.2  市场空间测算（TAM / SAM / SOM）", C_ABA)
    diagram(doc, "bp_funnel.png", "图 3  市场空间分层（SOM 数字为基于假设的目标）", width=12.5)
    callout(doc, "下文涉及的付费家庭数、客单价、营收等均为基于公开市场数据的合理假设，用于说明商业逻辑，最终以实际经营为准。", "tip")

    section(doc, "2.3  政策红利：从「鼓励」到「购买」", C_COACH)
    body(doc,
        "2024 年 7 月，中国残联、教育部、民政部、国家卫健委、国家医保局、共青团中央、全国妇联七部门联合印发"
        "《孤独症儿童关爱促进行动实施方案（2024—2028 年）》。其中「孤独症儿童家庭暖心行动」与我们的定位高度契合：")
    bullet(doc, "拓展「家长送训补贴」等救助内容，确保「应救尽救」；", C_COACH)
    bullet(doc, "推进政府购买服务，为家庭提供心理疏导、托养照料、喘息、社区支持等；", C_COACH)
    bullet(doc, "编写家长知识手册、普遍开展家长康复知识培训；", C_COACH)
    bullet(doc, "落实 29 项医疗康复项目纳入基本医保支付范围。", C_COACH)
    callout(doc,
        "政策正把「支持家长」从倡导变为可付费的公共服务采购项。这为我们的 G 端（政府购买）收入提供了明确依据，也是政府/公益申报的政策抓手。", "note")

    section(doc, "2.4  竞争格局：所有人都在「教孩子」", C_ABA)
    table(doc,
        headers=["代表机构", "模式", "融资 / 规模", "服务对象"],
        data=[
            ["大米和小米", "线下连锁 + RICE 干预体系", "融资 1.4 亿+，30+ 中心", "孩子"],
            ["东方启音", "语言康复切入连锁", "累计融资约 2.8 亿美元", "孩子"],
            ["恩启", "互联网 + 康复教育", "数千万元级融资", "孩子 / 康复师"],
            ["ALSOLIFE", "数字化评估干预 + AI", "AI 认知机器人，成本降至数元/节", "孩子"],
            ["五彩鹿", "线下连锁", "14+ 分支，未融资", "孩子"],
            ["本项目", "AI 包裹式家庭支撑（双 App）", "—", "孩子 + 家长（家庭）"],
        ],
        widths=[3, 4.5, 4.5, 4])
    callout(doc,
        "结论：头部玩家几乎全部聚焦「孩子的康复」，且以重资产线下连锁为主。"
        "我们以「家长」为差异化切入、以 AI 实现轻资产规模化，与现有机构是互补而非正面竞争——它们的客户正是我们的客户。", "tip")
    doc.add_page_break()

    # ---------- 三 ----------
    chapter_banner(doc, "三", "产品与解决方案", C_ABA)

    section(doc, "3.1  包裹式家庭支撑模型", C_ABA)
    body(doc,
        "我们用一个平台、两条线，把整个家庭稳稳「包裹」起来：「帮孩子」和「顾家长」并行，"
        "通过同账号、数据互通形成正向循环。")
    diagram(doc, "bp_solution.png", "图 4  包裹式家庭支撑：帮孩子 + 顾家长", width=15.0)

    section(doc, "3.2  ABA 智能助手（帮孩子）", C_ABA)
    body(doc, "面向孩子的科学干预核心工具，把专业 ABA 流程「翻译」成家长在家就能执行的标准动作。")
    bullet(doc, "122 项技能的入门评估与训练任务自动推荐；")
    bullet(doc, "DTT 离散试次标准记录（独立/语言/示范/辅助/错误四级辅助）；")
    bullet(doc, "AI 专业问答、自动训练计划、数据看板与结构化进展报告；")
    bullet(doc, "图片卡片等视觉教具，降低家庭执行门槛。")

    section(doc, "3.3  人生教练（顾家长）", C_COACH)
    body(doc, "面向家长本人的心理支撑系统，基于 ACT（接纳与承诺疗法）与正念心理学，是我们最大的差异化所在。")
    bullet(doc, "24 小时 AI 教练对话，随时倾诉、即时回应；", C_COACH)
    bullet(doc, "情绪追踪与心理状态识别，带安全层对危机信号进行兜底引导；", C_COACH)
    bullet(doc, "围绕焦虑/自责/疲惫/社交/亲子等议题的个性化成长路径（ACT 六阶段）；", C_COACH)
    bullet(doc, "覆盖 10 大主题的结构化知识库与正念练习。", C_COACH)

    section(doc, "3.4  技术与壁垒", C_ABA)
    table(doc,
        headers=["层面", "我们的能力 / 壁垒"],
        data=[
            ["产品壁垒", "国内少见的「孩子 + 家长」双轮一体化，已有可运行产品与数据闭环"],
            ["数据壁垒", "孩子干预数据 × 家长情绪数据的家庭级长周期数据，越用越准、难以复制"],
            ["算法壁垒", "ACT 引擎 + 安全层 + 个性化议题检测，将专业心理方法工程化"],
            ["成本结构", "AI 驱动的轻资产模式，单家庭边际服务成本极低，可规模化下沉"],
            ["合规安全", "数据本地化存储、危机识别与转介机制，契合公共服务采购要求"],
        ],
        widths=[3, 13])
    doc.add_page_break()

    # ---------- 四 ----------
    chapter_banner(doc, "四", "商业模式", C_ABA)

    section(doc, "4.1  三条收入线（C / B / G）", C_ABA)
    table(doc,
        headers=["收入线", "客户", "形态", "说明"],
        data=[
            ["C 端订阅", "自闭症儿童家庭", "会员订阅 / 增值功能", "家长直接付费，月/年费，含教练对话与高级报告"],
            ["B 端合作", "康复机构 / 医院", "SaaS 授权 + 分成", "机构把平台作为家庭作业与家长支持的延伸，按席位/家庭付费"],
            ["G 端购买", "残联 / 政府 / 公益", "政府购买服务 / 项目", "承接家长心理疏导、喘息、培训等公共服务采购"],
        ],
        widths=[2.5, 3.5, 4, 6])
    callout(doc,
        "三线协同：C 端验证需求与口碑，B 端借机构现有家庭快速放量，G 端用政策资金降低家庭付费门槛、扩大可及性。", "tip")

    section(doc, "4.2  定价假设", C_ABA)
    table(doc,
        headers=["产品", "定价（假设）", "对标 / 依据"],
        data=[
            ["家庭基础版", "免费", "获客入口，开放基础评估与记录"],
            ["家庭会员版", "约 49–99 元/月 或 599–999 元/年", "远低于线下单次干预数百元的成本"],
            ["机构 SaaS", "约 200–500 元/家庭/年", "按服务家庭数授权"],
            ["政府项目", "按项目打包", "依地方康复救助/购买服务标准"],
        ],
        widths=[3.5, 6, 6.5])

    section(doc, "4.3  获客渠道", C_ABA)
    bullet(doc, "B 端机构导流：与康复机构合作，把家长支持作为其服务的自然延伸；")
    bullet(doc, "G 端政策渠道：通过残联定点机构、政府购买项目触达家庭；")
    bullet(doc, "内容与社群：家长最信任「过来人」，以专业内容 + 家长社区形成口碑裂变；")
    bullet(doc, "医院 / 诊断机构转介：在确诊这一高需求时点切入。")
    doc.add_page_break()

    # ---------- 五 ----------
    chapter_banner(doc, "五", "运营与落地路径", C_ABA)
    diagram(doc, "bp_roadmap.png", "图 5  三步走落地路径", width=15.5)
    section(doc, "5.1  分阶段目标", C_ABA)
    table(doc,
        headers=["阶段", "重点", "关键目标（假设）"],
        data=[
            ["第 1 年 · 打磨验证", "双 App 数据闭环、跑通单城模型", "3 城试点，覆盖 1–2 万家庭，签约 2–3 家头部机构"],
            ["第 2 年 · 复制合作", "扩城 + 接入政府购买 + B 端标准化", "扩展至 6 城，落地 1–2 个政府购买项目"],
            ["第 3 年 · 规模壁垒", "全国铺开、数据/算法壁垒成型", "覆盖 15+ 城，建立家长社区生态"],
        ],
        widths=[4, 5, 7])
    section(doc, "5.2  政府 / 公益合作路径", C_COACH)
    bullet(doc, "申请成为残疾儿童康复救助定点服务机构的数字化服务合作方；", C_COACH)
    bullet(doc, "承接家长心理疏导、喘息服务、家长培训等政府购买项目；", C_COACH)
    bullet(doc, "与公益基金会联合开展面向低收入家庭的免费/补贴计划，兼顾社会价值与可及性。", C_COACH)
    doc.add_page_break()

    # ---------- 六 ----------
    chapter_banner(doc, "六", "财务预测（基于假设）", C_ABA)
    callout(doc, "以下数据均为基于公开市场信息的情景假设，用于说明增长与收入结构逻辑，非经营承诺。", "tip")
    diagram(doc, "bp_revenue.png", "图 6  三年营收预测与成熟期收入结构（假设）", width=15.5)
    section(doc, "6.1  三年营收预测（万元）", C_ABA)
    table(doc,
        headers=["收入线", "第 1 年", "第 2 年", "第 3 年"],
        data=[
            ["C 端家庭订阅", "180", "720", "2,100"],
            ["B 端机构合作", "120", "520", "1,500"],
            ["G 端政府购买", "200", "600", "1,600"],
            ["合计", "500", "1,840", "5,200"],
        ],
        widths=[5, 3.7, 3.7, 3.7])
    section(doc, "6.2  关键假设", C_ABA)
    bullet(doc, "付费家庭数：第 1 年约 0.3 万、第 3 年约 3–5 万（基于试点城市渗透率假设）；")
    bullet(doc, "C 端年客单价约 600 元，付费转化率随口碑逐年提升；")
    bullet(doc, "B 端按合作机构服务家庭数授权，G 端按地方购买服务项目打包。")
    doc.add_page_break()

    # ---------- 七 ----------
    chapter_banner(doc, "七", "融资计划", C_ABA)
    section(doc, "7.1  本轮融资", C_ABA)
    body(doc, "本轮拟引入战略 / 财务投资，用于产品打磨、试点城市落地与政府/机构合作拓展。具体融资额与估值待与投资方沟通确定。")
    table(doc,
        headers=["资金用途", "占比（建议）", "说明"],
        data=[
            ["产品与研发", "约 40%", "AI 引擎、安全层、数据平台与合规建设"],
            ["市场与渠道", "约 30%", "B 端机构合作、G 端政府项目、家长社群运营"],
            ["团队建设", "约 20%", "心理/ABA 专业团队、算法与产品人才"],
            ["运营储备", "约 10%", "试点城市落地与日常运营"],
        ],
        widths=[3.5, 3.5, 9])
    section(doc, "7.2  里程碑", C_ABA)
    numbered(doc, 1, "6 个月内：完成双 App 数据闭环，首批试点城市与机构签约；")
    numbered(doc, 2, "12 个月内：跑通单城单位经济模型，落地首个政府购买项目；")
    numbered(doc, 3, "24 个月内：扩展至 6 城，验证可复制性，启动下一轮融资。")
    doc.add_page_break()

    # ---------- 八 ----------
    chapter_banner(doc, "八", "团队与风险", C_ABA)
    section(doc, "8.1  团队（待补充）", C_ABA)
    table(doc,
        headers=["角色", "职责", "人选 / 背景"],
        data=[
            ["创始人 / CEO", "战略、融资、政府关系", "（待补充）"],
            ["产品 / 技术负责人", "双 App、AI 引擎与数据平台", "（待补充）"],
            ["心理 / ABA 专业负责人", "ACT 内容、循证与安全合规", "（待补充）"],
            ["市场 / BD 负责人", "机构合作与政府购买", "（待补充）"],
        ],
        widths=[4, 7, 5])
    section(doc, "8.2  风险与应对", C_ABA)
    table(doc,
        headers=["风险", "应对"],
        data=[
            ["AI 心理建议的安全与合规", "安全层 + 危机识别转介机制，明确不替代专业治疗的边界，引入专业督导"],
            ["家长付费意愿不确定", "免费基础版获客 + G 端补贴降低门槛 + B 端机构背书"],
            ["政策落地节奏不一", "C/B/G 三线并行，不依赖单一渠道"],
            ["数据隐私", "数据本地化存储与权限控制，契合公共服务采购的合规要求"],
        ],
        widths=[5, 11])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    section(doc, "数据来源", C_MUTE)
    for src in [
        "中国残联 2023 年残疾人普查报告；《孤独症儿童关爱促进行动实施方案（2024—2028 年）》（七部门）",
        "国家卫健委儿童孤独症患病率数据（2022）",
        "北大医疗脑健康《2020 年度儿童发展障碍康复报告》",
        "各机构融资 / 规模数据来自天眼查、企查查及公开媒体报道",
        "数字疗法市场规模等来自公开行业研究报道",
    ]:
        pp = doc.add_paragraph(); pp.paragraph_format.left_indent = Cm(0.5)
        pp.paragraph_format.space_after = Pt(2)
        font(pp.add_run("· "), 9, C_MUTE)
        font(pp.add_run(src), 9, C_MUTE)
    pe = doc.add_paragraph(); pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pe.paragraph_format.space_before = Pt(10)
    font(pe.add_run("— 本商业计划书含基于公开数据的假设性测算，仅供沟通参考 —"), 9, C_MUTE, italic=True)

    out = "/Users/wangxin/Documents/bjcontents/AI_codex/星星家庭_商业计划书.docx"
    doc.save(out)
    print("已生成：" + out)
    return out


if __name__ == "__main__":
    build()
