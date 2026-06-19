"""
生成 ABA 智能助手 · 用户手册 Word 文档（v2 排版）
- 颜色与产品封面统一：ABA 蓝 / Coach 青绿
- 章节使用色块编号 + 标题的"杂志式"样式
- 去掉正文标题里的 emoji，改用色条与色点引导阅读
- 表头统一 ABA 蓝，斑马行更柔和；提示框统一缩进与左侧色条
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================
# 颜色（与封面 / 产品 UI 一致）
# ============================================================
COLOR_ABA       = RGBColor(0x3A, 0x7B, 0xC8)  # ABA 主蓝
COLOR_ABA_DARK  = RGBColor(0x1F, 0x4E, 0x8C)  # 深蓝（标题强调）
COLOR_COACH     = RGBColor(0x0F, 0x76, 0x6E)  # 教练青绿
COLOR_COACH_DK  = RGBColor(0x0A, 0x4F, 0x4A)
COLOR_ACCENT    = RGBColor(0xE8, 0x7A, 0x4E)  # 暖橙（点缀，与封面圆点呼应）

# 中性
COLOR_INK       = RGBColor(0x1F, 0x29, 0x37)  # 正文深灰
COLOR_TEXT      = RGBColor(0x37, 0x41, 0x51)
COLOR_MUTE      = RGBColor(0x6B, 0x72, 0x80)
COLOR_LINE      = RGBColor(0xE5, 0xE7, 0xEB)
COLOR_ZEBRA     = RGBColor(0xF7, 0xFA, 0xFC)  # 极浅蓝灰
COLOR_TIP_BG    = RGBColor(0xEC, 0xF4, 0xFB)  # 提示框淡蓝
COLOR_WARN      = RGBColor(0xC2, 0x41, 0x0C)
COLOR_DANGER    = RGBColor(0xB9, 0x1C, 0x1C)
WHITE           = RGBColor(0xFF, 0xFF, 0xFF)

DIAGRAM_DIR = "/Users/wangxin/Documents/bjcontents/AI_codex/diagrams/"

# ============================================================
# XML 工具
# ============================================================

def _hex(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _hex(color))
    tcPr.append(shd)

def set_cell_borders(cell, color=COLOR_LINE, sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), _hex(color))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_shading(paragraph, color: RGBColor):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _hex(color))
    pPr.append(shd)

def set_para_left_border(paragraph, color: RGBColor, sz='24'):
    """段落左侧粗色条，营造引用块感。"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), _hex(color))
    pBdr.append(left)
    pPr.append(pBdr)

CJK_FONT = 'Noto Sans SC'

def set_run_font(run, size=11, color=COLOR_TEXT, bold=False, italic=False, name=CJK_FONT):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    # 统一中英文字体：ascii / hAnsi / eastAsia 全部指定，避免中文回退成宋体或忽黑忽宋
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)

# ============================================================
# 排版构件
# ============================================================

def chapter_banner(doc, num_cn, title, accent=COLOR_ABA):
    """章节大标题：左侧色块编号 + 右侧标题，下方分隔细线。"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(2.2)
    tbl.columns[1].width = Cm(13.5)

    # 左：色块编号
    c0 = tbl.rows[0].cells[0]
    c0.width = Cm(2.2)
    set_cell_bg(c0, accent)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(num_cn)
    set_run_font(r0, size=28, color=WHITE, bold=True)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 右：标题
    c1 = tbl.rows[0].cells[1]
    c1.width = Cm(13.5)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.left_indent = Cm(0.4)
    r1 = p1.add_run(title)
    set_run_font(r1, size=22, color=COLOR_INK, bold=True)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 去掉表格边框
    for cell in (c0, c1):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'nil')
            tcBorders.append(el)
        tcPr.append(tcBorders)

    # 下方细分隔线
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(10)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), _hex(accent))
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_heading(doc, text, accent=COLOR_ABA):
    """二级标题：左侧色点 + 标题。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r_dot = p.add_run("■  ")
    set_run_font(r_dot, size=12, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=14, color=COLOR_INK, bold=True)
    return p

def sub_heading(doc, text, accent=COLOR_ABA):
    """三级小标题。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=11.5, color=accent, bold=True)
    return p

def body(doc, text, size=10.5, leading=1.45):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = leading
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=size, color=COLOR_TEXT)
    return p

def bullet(doc, text, indent_cm=0.8, accent=COLOR_ABA):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    r_b = p.add_run("• ")
    set_run_font(r_b, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=COLOR_TEXT)
    return p

def numbered_step(doc, n, text, accent=COLOR_ABA):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    r_n = p.add_run(f"{n:>2}.  ")
    set_run_font(r_n, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=COLOR_TEXT)
    return p

def tip_box(doc, text, kind="tip"):
    """带左侧色条的提示框。"""
    if kind == "tip":
        label, accent = "小贴士", COLOR_ABA
    elif kind == "note":
        label, accent = "提示", COLOR_COACH
    elif kind == "warning":
        label, accent = "注意", COLOR_WARN
    elif kind == "danger":
        label, accent = "重要", COLOR_DANGER
    else:
        label, accent = "说明", COLOR_ABA

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.line_spacing = 1.4
    set_para_left_border(p, accent, sz='28')
    set_para_shading(p, COLOR_TIP_BG if kind in ('tip', 'note') else RGBColor(0xFD, 0xF2, 0xF2))

    r_lab = p.add_run(f" {label}  ")
    set_run_font(r_lab, size=9.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10, color=COLOR_TEXT)
    return p

def styled_table(doc, headers, data, col_widths=None, accent=COLOR_ABA):
    rows = len(data) + 1
    cols = len(headers)
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, accent)
        set_cell_borders(cell, color=accent)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h)
        set_run_font(r, size=10, color=WHITE, bold=True)

    # 数据行
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg = COLOR_ZEBRA if r_idx % 2 == 0 else WHITE
        for c_idx, txt in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, color=COLOR_LINE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(txt))
            set_run_font(r, size=9.5, color=COLOR_TEXT)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_diagram(doc, filename, caption="", width_cm=14.5):
    path = os.path.join(DIAGRAM_DIR, filename)
    if not os.path.exists(path):
        print(f"  缺图：{filename}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(10)
        r = p2.add_run(caption)
        set_run_font(r, size=9, color=COLOR_MUTE, italic=True)

# ============================================================
# 页脚页码
# ============================================================

def _page_field(paragraph):
    run = paragraph.add_run()
    a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
    b = OxmlElement('w:instrText'); b.set(qn('xml:space'), 'preserve'); b.text = 'PAGE'
    c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'end')
    run._r.append(a); run._r.append(b); run._r.append(c)
    set_run_font(run, size=9, color=COLOR_MUTE)
    return run

def setup_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("ABA 智能助手 · 用户手册      第 ")
    set_run_font(r1, size=9, color=COLOR_MUTE)
    _page_field(fp)
    r2 = fp.add_run(" 页")
    set_run_font(r2, size=9, color=COLOR_MUTE)

# ============================================================
# 主文档
# ============================================================

def _set_default_font(doc, name=CJK_FONT):
    style = doc.styles['Normal']
    style.font.name = name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)


def build_document():
    doc = Document()
    _set_default_font(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)

    setup_footer(doc)

    # -------- 封面（整页图） --------
    p_cover = doc.add_paragraph()
    p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cover.paragraph_format.space_before = Pt(0)
    p_cover.paragraph_format.space_after = Pt(0)
    _cover_path = os.path.join(DIAGRAM_DIR, "00_cover.png")
    if os.path.exists(_cover_path):
        p_cover.add_run().add_picture(_cover_path, width=Cm(13.5))
    else:
        r = p_cover.add_run("ABA 智能助手 · 用户手册")
        set_run_font(r, size=26, color=COLOR_ABA_DARK, bold=True)
    doc.add_page_break()

    # -------- 目录 --------
    p_toc = doc.add_paragraph()
    r_toc = p_toc.add_run("目  录")
    set_run_font(r_toc, size=20, color=COLOR_ABA_DARK, bold=True)
    p_toc.paragraph_format.space_after = Pt(4)

    p_sub = doc.add_paragraph()
    set_para_left_border(p_sub, COLOR_ABA, sz='24')
    r_sub = p_sub.add_run(" CONTENTS")
    set_run_font(r_sub, size=10, color=COLOR_MUTE, bold=True)
    p_sub.paragraph_format.space_after = Pt(14)

    toc_items = [
        ("一", "产品概述",         "了解 ABA 智能助手是什么、包含哪些模块"),
        ("二", "快速开始",         "注册登录、添加孩子档案、第一次评估"),
        ("三", "ABA 智能助手",     "主应用：问答、档案、训练、看板"),
        ("四", "人生教练",         "面向家长：情绪、成长路径、知识库"),
        ("五", "常见问题 FAQ",     "使用过程中可能遇到的问题及解答"),
        ("六", "技术支持",         "紧急联系方式、产品信息"),
    ]
    for num, title, desc in toc_items:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.autofit = False
        tbl.columns[0].width = Cm(1.2)
        tbl.columns[1].width = Cm(5.5)
        tbl.columns[2].width = Cm(9.0)

        c0 = tbl.rows[0].cells[0]; c0.width = Cm(1.2)
        set_cell_bg(c0, COLOR_ABA); c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(num); set_run_font(r0, size=14, color=WHITE, bold=True)

        c1 = tbl.rows[0].cells[1]; c1.width = Cm(5.5)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        p1.paragraph_format.left_indent = Cm(0.3)
        r1 = p1.add_run(title); set_run_font(r1, size=12.5, color=COLOR_INK, bold=True)

        c2 = tbl.rows[0].cells[2]; c2.width = Cm(9.0)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(desc); set_run_font(r2, size=10, color=COLOR_MUTE)

        for cell in tbl.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'right'):
                el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'), 'nil'); tcBorders.append(el)
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), _hex(COLOR_LINE))
            tcBorders.append(bottom)
            tcPr.append(tcBorders)

        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    doc.add_page_break()

    # =========================================================
    # 一、产品概述
    # =========================================================
    chapter_banner(doc, "一", "产品概述", COLOR_ABA)

    section_heading(doc, "1.1  这是什么？", COLOR_ABA)
    body(doc,
         "ABA 智能助手是一款面向自闭症儿童家长的 AI 辅助工具，帮助你在家庭环境中科学地进行 "
         "ABA（应用行为分析）干预训练，同时照顾好你自己的心理健康。")

    section_heading(doc, "1.2  两大模块", COLOR_ABA)
    styled_table(doc,
        headers=["模块", "访问端口", "服务对象", "核心功能"],
        data=[
            ["ABA 智能助手（主应用）", ":8501", "孩子", "ABA 问答、训练记录、技能评估、任务规划、数据看板"],
            ["人生教练（独立应用）",   ":8503", "家长本人", "情绪支持、心理成长、知识学习、个人记录、成长路径"],
        ],
        col_widths=[4.5, 2.4, 2.4, 6.7])

    tip_box(doc,
        "两个模块共享同一个账号系统。孩子在主应用中的档案，也会被人生教练读取，用于提供个性化建议。",
        "tip")

    add_diagram(doc, "01_product_overview.png", "图 1  ABA 智能助手产品结构示意")

    section_heading(doc, "1.3  使用前准备", COLOR_ABA)
    sub_heading(doc, "系统要求", COLOR_ABA)
    bullet(doc, "支持现代浏览器的电脑（Chrome / Edge / Firefox 均可）")
    bullet(doc, "或直接在 WorkBuddy 小程序中访问")

    sub_heading(doc, "第一次使用（推荐步骤）", COLOR_ABA)
    numbered_step(doc, 1, "在 ABA 主应用（端口 8501）注册一个账号")
    numbered_step(doc, 2, "登录后，先添加孩子档案（填写孩子基本信息）")
    numbered_step(doc, 3, "带孩子做一次「入门评估」，系统会推荐训练起点")
    numbered_step(doc, 4, "根据推荐开始日常训练")

    doc.add_page_break()

    # =========================================================
    # 二、快速开始
    # =========================================================
    chapter_banner(doc, "二", "快速开始", COLOR_ABA)
    add_diagram(doc, "02_quick_start.png", "图 2  新用户 5 步快速上手流程")

    section_heading(doc, "2.1  注册与登录", COLOR_ABA)
    sub_heading(doc, "注册账号", COLOR_ABA)
    for i, s in enumerate([
        "打开 ABA 主应用（默认地址 http://localhost:8501）",
        "在左侧边栏选择「注册」标签页",
        "输入用户名（建议使用孩子小名）和密码",
        "点击「注册」按钮",
        "注册成功后自动跳转至登录状态",
    ]):
        numbered_step(doc, i+1, s)

    sub_heading(doc, "登录", COLOR_ABA)
    body(doc, "每次打开应用，在左侧边栏输入用户名和密码，点击「登录」即可。")

    section_heading(doc, "2.2  添加第一个孩子档案", COLOR_ABA)
    for i, (t, d) in enumerate([
        ("进入孩子档案页面", "左侧边栏点击「孩子档案」"),
        ("切换到添加档案",   "点击「添加档案」标签页"),
        ("填写信息",         "按提示填写下方表格中的字段"),
        ("保存档案",         "点击「添加档案」保存"),
    ]):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(2)
        r_s = p.add_run(f"Step {i+1}   ")
        set_run_font(r_s, size=10.5, color=COLOR_ABA, bold=True)
        r_t = p.add_run(t + "   ")
        set_run_font(r_t, size=10.5, color=COLOR_INK, bold=True)
        r_d = p.add_run(d)
        set_run_font(r_d, size=10, color=COLOR_MUTE)

    sub_heading(doc, "档案字段说明", COLOR_ABA)
    styled_table(doc,
        headers=["字段", "是否必填", "说明"],
        data=[
            ["孩子姓名", "必填", "可以用小名"],
            ["出生日期", "选填", "用于自动计算年龄"],
            ["诊断情况", "选填", "自闭症 / 发育迟缓 / 感统失调 / 语言发育迟缓 / 其他"],
            ["诊断时间", "选填", "帮助追踪干预时长"],
            ["干预目标", "选填", "描述当前最想提升的能力，如「表达需求」「安坐注意力」"],
            ["备注",     "选填", "任何想补充的信息"],
        ],
        col_widths=[3.5, 2.5, 10])

    section_heading(doc, "2.3  入门评估（可选但推荐）", COLOR_ABA)
    body(doc,
        "入门评估能帮你快速了解孩子各能力领域的当前水平，并自动推荐合适的训练起点。"
        "评估不需要准备任何材料，凭日常观察回答即可。")

    for i, (t, d) in enumerate([
        ("选择入口",   "左侧边栏点击「入门评估」，选择孩子"),
        ("回答问题",   "逐题回答 27 道是/否题（覆盖 9 个能力领域），约 5 分钟"),
        ("查看结果",   "每个领域的得分百分比与掌握状态：已掌握 / 部分掌握 / 建议优先训练"),
        ("生成任务",   "点击「一键生成全部训练任务」，将推荐技能加入任务清单"),
    ]):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(2)
        r_s = p.add_run(f"Step {i+1}   ")
        set_run_font(r_s, size=10.5, color=COLOR_ABA, bold=True)
        r_t = p.add_run(t + "   ")
        set_run_font(r_t, size=10.5, color=COLOR_INK, bold=True)
        r_d = p.add_run(d)
        set_run_font(r_d, size=10, color=COLOR_MUTE)

    tip_box(doc, "建议每季度重做一次入门评估，便于横向对比孩子的进步。", "tip")

    doc.add_page_break()

    # =========================================================
    # 三、ABA 智能助手
    # =========================================================
    chapter_banner(doc, "三", "ABA 智能助手（主应用）", COLOR_ABA)

    body(doc,
        "ABA 智能助手是你日常干预训练的核心工具，包含 10 大功能模块，"
        "帮助你在家庭环境中科学、系统地开展 ABA 训练。")

    section_heading(doc, "3.1  功能导航", COLOR_ABA)
    styled_table(doc,
        headers=["功能入口", "核心用途", "推荐频率"],
        data=[
            ["AI 助手",     "ABA 专业问答、日常咨询",          "每天"],
            ["孩子档案",     "管理孩子信息、查看干预目标",       "随时"],
            ["入门评估",     "了解孩子能力水平、生成训练建议",   "初次 + 每季度"],
            ["任务清单",     "查看和管理训练计划",               "每天"],
            ["训练记录",     "按 DTT 标准流程记录每次训练",      "每次训练后"],
            ["图片卡片",     "ABA 视觉教具展示和互动练习",       "训练时"],
            ["进展记录",     "查看技能掌握趋势和时间线",         "每周"],
            ["数据看板",     "图表化展示训练全貌",               "每周"],
            ["报告中心",     "生成结构化训练总结报告",           "每月"],
            ["进入人生教练", "跳转到人生教练模块",               "需要时"],
        ],
        col_widths=[3.5, 8, 4])

    section_heading(doc, "3.2  AI 助手（对话问答）", COLOR_ABA)
    body(doc,
        "这是你的日常咨询入口。任何时候遇到 ABA 相关问题，都可以来这里提问。"
        "AI 会结合专业知识给出温暖、实用的回答。")

    sub_heading(doc, "可以问什么？", COLOR_ABA)
    styled_table(doc,
        headers=["问题类型", "示例问题"],
        data=[
            ["ABA 概念", "什么是 DTT？强化和惩罚有什么区别？"],
            ["行为问题", "孩子总是哭闹怎么办？如何减少自伤行为？"],
            ["训练技巧", "怎么教孩子提要求？DTT 的具体步骤是什么？"],
            ["目标设定", "3 岁孩子应该练什么？如何制定 SMART 目标？"],
            ["数据记录", "如何计算正确率？什么叫辅助递减？"],
        ],
        col_widths=[3.5, 12])

    tip_box(doc, "AI 支持连续对话，你可以接着上一轮的回答继续问，AI 会记住前文上下文。", "tip")

    section_heading(doc, "3.3  训练记录（DTT 标准流程）", COLOR_ABA)
    body(doc,
        "DTT（离散试次训练）是 ABA 中最基础的教学方法。系统会引导你按照标准流程，"
        "记录每一次试训的结果。")

    sub_heading(doc, "ABA 四级辅助记录体系", COLOR_ABA)
    styled_table(doc,
        headers=["按钮", "含义", "ABA 术语"],
        data=[
            ["独立", "孩子无需任何辅助，独立完成",                  "Independent (I)"],
            ["语言", "需要语言提示后正确（如「你说想要什么？」）",   "Verbal prompt (V)"],
            ["示范", "需要示范动作后正确（如家长先说一遍）",         "Model prompt (M)"],
            ["辅助", "需要身体辅助后正确（手把手）",                 "Physical assist (P)"],
            ["错误", "孩子回答错误，或无反应",                       "Error (E)"],
        ],
        col_widths=[2.5, 8.5, 4.5])

    add_diagram(doc, "03_dtt_levels.png", "图 3  DTT 四级辅助记录体系可视化")

    sub_heading(doc, "正确率计算规则", COLOR_ABA)
    body(doc, "正确率 = 独立正确次数 ÷ 总试次数 × 100%。只有完全独立的反应才算「掌握」（这是 ABA 标准做法）。")

    section_heading(doc, "3.4  任务清单", COLOR_ABA)
    body(doc,
        "任务清单是你为孩子规划的「训练计划」。系统会根据孩子的年龄、掌握情况，"
        "从课程技能树（122 项技能）中自动推荐合适的训练技能。")

    styled_table(doc,
        headers=["标签页", "功能说明"],
        data=[
            ["当前任务", "默认显示正在训练的技能卡片，包含进度条与操作按钮"],
            ["已完成",   "查看已经掌握或手动完成的任务，可重新加入或删除"],
            ["手动添加", "除了系统自动生成，你也可以手动添加任务"],
        ],
        col_widths=[3.5, 12])

    section_heading(doc, "3.5  数据看板", COLOR_ABA)
    body(doc, "数据看板用图表形式直观展示孩子的训练全貌，帮助你快速判断训练效果。")
    styled_table(doc,
        headers=["核心指标", "说明"],
        data=[
            ["训练次数",       "已完成的所有训练 session 数"],
            ["本周训练",       "最近 7 天内的训练次数"],
            ["训练技能数",     "不重复计数，共练过几个不同技能"],
            ["已掌握技能数",   "连续 3 次 session 正确率 ≥ 80% 的技能数"],
            ["总独立正确率",   "所有训练中「独立正确」次数的占比"],
        ],
        col_widths=[4, 11.5])

    doc.add_page_break()

    # =========================================================
    # 四、人生教练
    # =========================================================
    chapter_banner(doc, "四", "人生教练（面向家长）", COLOR_COACH)

    body(doc,
        "人生教练专注于服务家长本人的心理健康和个人成长。基于 ACT（接纳与承诺疗法）"
        "和正念心理学，为你提供温暖的心理支持和实用的应对建议。")

    tip_box(doc,
        "如果你已经在 ABA 主应用登录，点击「进入人生教练」会自动登录，无需重新输入密码（SSO 单点登录）。",
        "note")

    section_heading(doc, "4.1  ACT 六大核心过程", COLOR_COACH)
    body(doc, "成长路径的理论基础，是人生教练所有功能背后的心理学框架。")
    styled_table(doc,
        headers=["阶段", "名称", "核心问题"],
        data=[
            ["1", "觉察与接纳",   "我能看到自己的情绪吗？能接纳它们的存在吗？"],
            ["2", "与想法解离",   "我能看到「想法 ≠ 事实」吗？不和消极想法认同吗？"],
            ["3", "当下觉察",     "我能回到此时此地吗？不被过去与未来拉着走？"],
            ["4", "看见背后的我", "除了「自闭症孩子的家长」，我还是谁？"],
            ["5", "价值观探索",   "对我真正重要的是什么？我想成为什么样的人？"],
            ["6", "承诺行动",     "我今天可以采取什么和价值观一致的小行动？"],
        ],
        col_widths=[1.5, 4, 10],
        accent=COLOR_COACH)

    add_diagram(doc, "04_act_process.png", "图 4  ACT 六大核心过程环形图")

    section_heading(doc, "4.2  首页（仪表盘）", COLOR_COACH)
    body(doc, "首页是你的「仪表盘」，快速总览今天的状况并提供快捷入口。")
    for feat, desc in [
        ("动态欢迎语",     "根据当前时间显示：早上好/下午好/晚上好 + 你的用户名"),
        ("成长足迹统计",   "三张数据卡片：情绪记录数、任务完成数、个人记录数"),
        ("每日教练建议",   "每天刷新的一段温暖建议，基于 ACT 和正念心理学"),
        ("快捷入口",       "8 个情绪按钮 + 6 个生活分类按钮，点击后直接进入对应功能"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r_b = p.add_run("• ")
        set_run_font(r_b, size=10.5, color=COLOR_COACH, bold=True)
        r_t = p.add_run(feat + "   ")
        set_run_font(r_t, size=10.5, color=COLOR_INK, bold=True)
        r_d = p.add_run(desc)
        set_run_font(r_d, size=10, color=COLOR_TEXT)

    section_heading(doc, "4.3  教练对话", COLOR_COACH)
    body(doc,
        "教练对话是人生教练的核心功能。你可以随时来这里倾诉、询问，"
        "获得温暖的心理支持和实用的应对建议。")
    sub_heading(doc, "情绪快捷按钮", COLOR_COACH)
    body(doc,
        "对话界面顶部提供情绪快捷按钮。点击任一按钮，会自动将对应情绪文字发送到对话中，"
        "教练立即回应。例如点击「焦虑」，相当于发送了「我现在感到焦虑」。")

    section_heading(doc, "4.4  情绪追踪", COLOR_COACH)
    body(doc, "系统记录自己的情绪模式，看见进步，并提前发现需要关注的心理状态。")
    styled_table(doc,
        headers=["字段", "是否必填", "说明"],
        data=[
            ["心情",                "必填", "11 种选项：很好 / 平静 / 一般 / 不好 / 焦虑 等"],
            ["发生了什么（触发事件）", "选填", "记录触发这次情绪的具体事件"],
            ["情绪强度",            "选填", "1（很轻微）～ 10（非常强烈），滑块选择"],
            ["身体的感受",          "选填", "情绪在身体中的表现，如「胸口闷」「肩膀紧绷」"],
            ["当时的想法",          "选填", "记录自动思维，如「别人都在看我们」"],
            ["想补充的话",          "选填", "任何想补充的内容"],
        ],
        col_widths=[4, 2.5, 9.5],
        accent=COLOR_COACH)

    section_heading(doc, "4.5  成长路径（核心系统）", COLOR_COACH)
    body(doc,
        "成长路径基于 ACT 六大核心过程，围绕一个具体的个人议题，带你走完一段有结构的心理成长旅程。"
        "系统会根据你提出的议题自动检测类型，并生成个性化的练习路径。")

    styled_table(doc,
        headers=["议题类型", "关键词示例", "个性化内容"],
        data=[
            ["焦虑类",       "焦虑、担心、害怕、紧张",              "焦虑身体地图、4-7-8 呼吸法、焦虑冲浪、安全锚"],
            ["自责类",       "自责、内疚、我的错、对不起",          "给批评家取名字、慈悲朋友视角、写慈悲信"],
            ["疲惫类",       "累、疲惫、撑不住、耗竭",              "精力地图、氧气面罩法则、微型休息菜单"],
            ["社交类",       "社交、孤立、一个人、别人的眼光",      "聚光灯效应检验、暴露练习、找到同类社群"],
            ["亲子关系类",   "孩子、亲子关系、沟通、不听话",        "情绪触发链、暂停-呼吸-回应、正念亲子游戏"],
            ["通用模板",     "不匹配以上任何类型",                  "固定 ACT 六阶段（28 个通用练习）"],
        ],
        col_widths=[3, 5, 7.5],
        accent=COLOR_COACH)

    add_diagram(doc, "05_growth_path.png", "图 5  成长路径个性化议题驱动流程")

    sub_heading(doc, "如何创建一个成长议题", COLOR_COACH)
    for i, s in enumerate([
        "在「已完成」或「暂停中」区域下方，找到「描述你的议题」输入框",
        "用你自己的话描述当前最困扰你的一件事",
        "点击「开启成长之旅」",
        "系统自动检测议题类型，生成专属的 6 阶段练习任务",
        "按阶段依次完成，循序渐进",
    ]):
        numbered_step(doc, i+1, s, accent=COLOR_COACH)

    section_heading(doc, "4.6  知识库", COLOR_COACH)
    body(doc, "结构化自助学习系统，覆盖 10 大主题领域。")
    styled_table(doc,
        headers=["#", "领域", "内容示例"],
        data=[
            ["1",  "核心理论",         "ACT 六大过程详解、正念基础、心理灵活性"],
            ["2",  "情绪管理",         "焦虑/悲伤/愤怒/自责/疲惫/孤独等情绪的应对策略"],
            ["3",  "关系与沟通",       "伴侣关系、家庭沟通、朋友社交、边界设定"],
            ["4",  "工作与平衡",       "工作与育儿平衡、职业发展、经济压力"],
            ["5",  "自我关怀",         "睡眠管理、运动、饮食、身体觉察、医疗健康"],
            ["6",  "时间管理",         "精力管理、习惯养成、日常仪式"],
            ["7",  "决策力与行动力",   "决策框架、拖延应对、适应与过渡"],
            ["8",  "自闭症养育专题",   "养育情绪、场景应对、支持网络构建"],
            ["9",  "正念练习库",       "各级别的结构化正念练习（1 分钟～30 分钟）"],
            ["10", "附加资源",         "推荐书籍、有用 App、紧急联系方式"],
        ],
        col_widths=[1, 4.5, 10],
        accent=COLOR_COACH)

    doc.add_page_break()

    # =========================================================
    # 五、FAQ
    # =========================================================
    chapter_banner(doc, "五", "常见问题 FAQ", COLOR_ABA)

    faq_sections = [
        ("账号与数据", COLOR_ABA, [
            ("我的数据安全吗？",
             "所有数据存储在你的本地服务器上，不会上传到任何第三方。只有你本人和授权的管理员可以访问。"),
            ("我可以导出我的数据吗？",
             "可以。请联系管理员，在专家后台中使用「导出数据」功能。"),
            ("忘记密码怎么办？",
             "目前需要联系管理员重置密码。建议在注册时记下密码或存入密码管理工具中。"),
        ]),
        ("ABA 训练", COLOR_ABA, [
            ("DTT 是什么？我从来没学过 ABA",
             "DTT（Discrete Trial Teaching，离散试次训练）是 ABA 中最基础的教学方法。"
             "简单说就是：发指令 → 等孩子反应 → 正确就强化，错误就纠错。系统会在训练中逐步引导你。"),
            ("正确率怎么算？",
             "只计算「独立正确」的次数。比如你记录了 10 次试次，其中 6 次是孩子独立正确的，"
             "那么正确率 = 60%。"),
            ("什么叫「掌握」？",
             "连续 3 次训练 session，每次的独立正确率都 ≥ 80%，系统就认为孩子已经掌握这个技能。"),
        ]),
        ("人生教练", COLOR_COACH, [
            ("人生教练能替代心理治疗吗？",
             "不能。人生教练提供日常心理支持和自助练习，但不能替代专业心理治疗。"
             "如果你有深度的心理困扰，请一定要寻求专业心理咨询师的帮助。"),
            ("成长路径一定要按阶段顺序走吗？",
             "是的，ACT 六大过程是有顺序的——先觉察，再解离，再当下，再价值观，最后行动。跳级可能会效果打折。"),
            ("我可以同时开多个成长议题吗？",
             "可以。先暂停当前议题，再开启新议题。但一次只能激活一个（即只能在一个议题中记录任务完成）。"),
        ]),
    ]

    for section_title, accent, qas in faq_sections:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(10)
        p_sec.paragraph_format.space_after = Pt(4)
        set_para_left_border(p_sec, accent, sz='28')
        r_sec = p_sec.add_run(f" {section_title}")
        set_run_font(r_sec, size=12, color=accent, bold=True)

        for q, a in qas:
            p_q = doc.add_paragraph()
            p_q.paragraph_format.left_indent = Cm(0.4)
            p_q.paragraph_format.space_before = Pt(4)
            p_q.paragraph_format.space_after = Pt(1)
            r_q1 = p_q.add_run("Q  ")
            set_run_font(r_q1, size=10.5, color=accent, bold=True)
            r_q2 = p_q.add_run(q)
            set_run_font(r_q2, size=10.5, color=COLOR_INK, bold=True)

            p_a = doc.add_paragraph()
            p_a.paragraph_format.left_indent = Cm(0.4)
            p_a.paragraph_format.space_after = Pt(3)
            p_a.paragraph_format.line_spacing = 1.4
            r_a1 = p_a.add_run("A  ")
            set_run_font(r_a1, size=10, color=COLOR_MUTE, bold=True)
            r_a2 = p_a.add_run(a)
            set_run_font(r_a2, size=10, color=COLOR_TEXT)

    doc.add_page_break()

    # =========================================================
    # 六、技术支持
    # =========================================================
    chapter_banner(doc, "六", "技术支持", COLOR_ABA)

    section_heading(doc, "6.1  产品信息", COLOR_ABA)
    styled_table(doc,
        headers=["项目", "内容"],
        data=[
            ["产品名称", "ABA 智能助手（含人生教练模块）"],
            ["当前版本", "v1.4.0"],
            ["最后更新", "2026-05-31"],
            ["技术支持", "请联系产品开发团队"],
        ],
        col_widths=[4, 12])

    section_heading(doc, "6.2  紧急联系方式（心理危机情况）", COLOR_DANGER)
    body(doc, "如果你或你认识的人正处于危险中，请立即拨打以下任一热线，或前往最近的医院急诊室。")
    for name, number in [
        ("全国 24 小时心理援助热线", "400-161-9995"),
        ("北京心理危机研究与干预中心", "010-82951332"),
        ("生命热线",                  "400-821-1215"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r_b = p.add_run("● ")
        set_run_font(r_b, size=10.5, color=COLOR_DANGER, bold=True)
        r_n = p.add_run(f"{name}    ")
        set_run_font(r_n, size=10.5, color=COLOR_INK, bold=True)
        r_num = p.add_run(number)
        set_run_font(r_num, size=11, color=COLOR_DANGER, bold=True)

    tip_box(doc,
        "本手册中的人生教练功能不能替代专业心理治疗。请在危机时刻第一时间联系上述专业资源。",
        "danger")

    # 尾注
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("— 本手册基于 v1.4.0 版本编写，功能更新后请以应用界面为准 —")
    set_run_font(r_end, size=9, color=COLOR_MUTE, italic=True)

    out = "/Users/wangxin/Documents/bjcontents/AI_codex/ABA智能助手_用户手册.docx"
    doc.save(out)
    print(f"已生成：{out}")
    return out

if __name__ == "__main__":
    build_document()
